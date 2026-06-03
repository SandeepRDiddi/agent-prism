from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Agent_Prism_End_to_End_Demo_Runbook.docx"


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
GREEN = "DFF3E8"
AMBER = "FFF2CC"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(9)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)


def set_table_borders(table, color="D7DEE8"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.first_child_found_in("w:tblCellMar")
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tbl_cell_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def make_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    set_cell_margins(table)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade(cell, LIGHT_BLUE)
        set_cell_text(cell, header, bold=True, color=INK)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            set_cell_text(cells[idx], str(text))
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    return table


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.15)
    r = p.add_run(text)
    r.font.name = "JetBrains Mono"
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string("243B61")
    shade_paragraph(p, "EEF3FA")
    return p


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_callout(doc, title, body, fill=CALLOUT):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    shade_paragraph(p, fill)
    r = p.add_run(f"{title}: ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    p.add_run(body)


def add_step(doc, number, title, body, command=None):
    p = doc.add_paragraph(style="Heading 3")
    p.add_run(f"Step {number}: {title}")
    for line in body:
        doc.add_paragraph(line, style="List Bullet")
    if command:
        add_code(doc, command)


def setup_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string("1F2937")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.text = "Agent Prism Demo Runbook"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = RGBColor.from_string("6B7280")

    footer = section.footer.paragraphs[0]
    footer.text = "Confidential demo guide | Agent Prism"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor.from_string("6B7280")


def build():
    doc = Document()
    setup_styles(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("Agent Prism End-to-End Enterprise Demo Runbook")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor.from_string(INK)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    r = subtitle.add_run("Postman API flow, SaaS login, Postgres verification, tenant isolation proof, and AI Advisor demo.")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string("4B5563")

    make_table(
        doc,
        ["Demo owner", "Environment", "Primary URL", "Artifact"],
        [["Sandeep / Agent Prism", "Render SaaS + Postgres", "https://agent-prism.onrender.com", "Postman + Browser + psql"]],
        [1.45, 1.65, 2.2, 1.2],
    )

    add_callout(
        doc,
        "Demo objective",
        "Show that Agent Prism is an enterprise-ready SaaS control plane: users log in by tenant, agents send telemetry through tenant API keys, data lands in Postgres, audit logs prove actions, and AI Advisor produces recommendations from tenant-scoped telemetry.",
        GREEN,
    )

    doc.add_heading("1. Demo Storyline", level=1)
    for item in [
        "A company owner logs in to the shared Agent Prism URL.",
        "The tenant creates API keys for agents and automation.",
        "Postman simulates agent traffic, gateway traffic, and connector events.",
        "Agent Prism shows dashboard, Token Coach, AI Advisor, audit logs, and API key governance.",
        "Postgres checks prove tenant-scoped storage and revocation behavior.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("2. Required Assets", level=1)
    make_table(
        doc,
        ["Asset", "Where it comes from", "Used for"],
        [
            ["Postman collection", "postman/agent-prism.postman_collection.json", "All API demo calls"],
            ["Postman environment", "postman/agent-prism.postman_environment.json", "base_url, admin_secret, api_key, key_id variables"],
            ["Owner email", "select email from users where role='owner';", "Human dashboard login"],
            ["Owner password", "Set via /api/admin/users/password", "Human dashboard login"],
            ["Tenant API key", "Admin API key endpoint or Admin tab", "Agent/Postman requests"],
            ["External DB URL", "Render Postgres External tab", "Local psql verification"],
        ],
        [1.6, 2.5, 2.2],
    )

    doc.add_heading("3. One-Time Postgres Preparation", level=1)
    add_step(
        doc,
        1,
        "Run login migration",
        ["Use the External Database URL from your laptop. The Internal URL works only inside Render."],
        'cd /Users/sandeepdiddi/Documents/agent-prism/agent-prism\npsql "$DATABASE_URL" -f db/migrations/001_dashboard_login.sql',
    )
    add_step(
        doc,
        2,
        "Find the owner email",
        ["This is the email entered during the first bootstrap."],
        'psql "$DATABASE_URL" -c "select id, tenant_id, email, name, role from users;"',
    )
    add_step(
        doc,
        3,
        "Set owner login password",
        ["Choose a demo password. This is the Agent Prism login password, not the DB password and not OpenRouter."],
        'export ACP_ADMIN_SECRET="your_render_admin_secret"\n\ncurl -s -X POST https://agent-prism.onrender.com/api/admin/users/password \\\n  -H "content-type: application/json" \\\n  -H "x-admin-secret: $ACP_ADMIN_SECRET" \\\n  -d \'{"email":"owner@example.com","password":"ChooseStrongPassword123"}\'',
    )
    add_callout(doc, "Expected result", 'The response returns a "user" object. If it says "User not found", replace owner@example.com with the actual email from the users table.', AMBER)

    doc.add_heading("4. Postman Import and Variables", level=1)
    doc.add_paragraph("Import both JSON files into Postman, then choose the Agent Prism Render environment.")
    make_table(
        doc,
        ["Variable", "Example value", "Notes"],
        [
            ["base_url", "https://agent-prism.onrender.com", "Shared SaaS URL"],
            ["admin_secret", "ACP_ADMIN_SECRET value", "Only for admin setup/recovery calls"],
            ["api_key", "acp_...", "Tenant API key for agent/API traffic"],
            ["key_id", "key_...", "Only needed when revoking a test key"],
            ["session_id", "sess_...", "Only needed for session update demo"],
        ],
        [1.5, 2.4, 2.6],
    )

    doc.add_page_break()
    doc.add_heading("5. Live Demo Run Order", level=1)
    demo_steps = [
        ("Health", "01 Setup and Admin / Health", "Confirm the service is up and using the expected storage backend."),
        ("Bootstrap Status", "01 Setup and Admin / Bootstrap Status", "Confirm tenant already exists."),
        ("Admin Create Browser API Key", "01 Setup and Admin / Admin Create Browser API Key", "Create an acp_ key if you need one for Postman."),
        ("Tenant Summary", "02 Tenant Dashboard / Tenant Summary", "Prove the API key maps to one tenant."),
        ("Generic Ingest", "05 Ingest Telemetry / Generic Ingest", "Create one synthetic agent run."),
        ("Dashboard Snapshot", "02 Tenant Dashboard / Dashboard Snapshot", "Show metrics now include the Postman run."),
        ("AI Advisor", "02 Tenant Dashboard / AI Advisor", "Show model-generated recommendations for the tenant telemetry."),
        ("Audit Logs", "02 Tenant Dashboard / Audit Logs", "Show key creation and user actions are recorded."),
        ("Create Tenant API Key", "03 API Keys / Create Tenant API Key", "Create Demo key B for revocation proof."),
        ("Revoke Tenant API Key", "03 API Keys / Revoke Tenant API Key", "Show lifecycle governance."),
    ]
    make_table(doc, ["Order", "Postman request", "Talk track"], [(i + 1, call, why) for i, (name, call, why) in enumerate(demo_steps)], [0.55, 2.5, 3.45])

    doc.add_heading("6. Browser Login Demo", level=1)
    add_step(
        doc,
        1,
        "Clear old machine key",
        ["This proves humans are not relying on localStorage API keys."],
        'localStorage.removeItem("acp_api_key")\nlocation.reload()',
    )
    add_step(
        doc,
        2,
        "Sign in",
        ["Use owner email and password. Agent Prism will set a secure HttpOnly session cookie named aps_session."],
        "Email: owner email from Postgres\nPassword: password set through admin endpoint",
    )
    add_step(
        doc,
        3,
        "Verify session",
        ["Run this in the browser console. authType should be session."],
        'fetch("/api/me").then(r => r.json()).then(console.log)',
    )

    doc.add_heading("7. Tenant Isolation and Revocation Proof", level=1)
    add_callout(
        doc,
        "Positioning",
        "With one tenant, the demo proves tenant-bound API keys, revocation, auditability, and DB tenant_id scoping. A second-tenant admin flow and Postgres Row Level Security are the next hardening step for cross-company proof.",
        CALLOUT,
    )
    for n, (title_text, body_text, cmd) in enumerate(
        [
            (
                "Create Demo key A and Demo key B",
                ["Use Postman Create Tenant API Key twice and save both returned acp_ keys."],
                'export KEY_A="acp_demo_key_a"\nexport KEY_B="acp_demo_key_b"',
            ),
            (
                "Both keys can read the same tenant",
                ["Both keys belong to the same tenant, so both should return the same tenant object."],
                'curl -s "$PRISM_URL/api/tenant" -H "x-api-key: $KEY_A"\ncurl -s "$PRISM_URL/api/tenant" -H "x-api-key: $KEY_B"',
            ),
            (
                "Write telemetry with Key A",
                ["Use Generic Ingest in Postman or this curl command."],
                'curl -s -X POST "$PRISM_URL/api/ingest" \\\n  -H "content-type: application/json" \\\n  -H "x-api-key: $KEY_A" \\\n  -d \'{"source":"generic","payload":{"agentName":"Tenant Isolation Test Agent","provider":"OpenAI","model":"gpt-4o-mini","taskType":"tenant-isolation-test","status":"success","latencyMs":1200,"tokensIn":1000,"tokensOut":200,"costUsd":0.002,"budgetUsd":0.01,"autonomyLevel":2,"retryCount":0,"toolCalls":1,"policyViolations":0,"userSatisfaction":5,"environment":"enterprise-demo","workflow":"tenant-isolation","team":"security"}}\'',
            ),
            (
                "Read telemetry with Key B",
                ["Key B sees the run because it belongs to the same tenant, not because data is global."],
                'curl -s "$PRISM_URL/api/runs" -H "x-api-key: $KEY_B"',
            ),
            (
                "Revoke Key B",
                ["List API keys, copy Demo key B id, set key_id in Postman, then run Revoke Tenant API Key."],
                'curl -s "$PRISM_URL/api/tenant/api-keys" -H "x-api-key: $KEY_A"\ncurl -s -X DELETE "$PRISM_URL/api/tenant/api-keys/$KEY_B_ID" -H "x-api-key: $KEY_A"',
            ),
            (
                "Confirm revoked key fails",
                ["Expected result is HTTP 401 unauthorized."],
                'curl -i "$PRISM_URL/api/tenant" -H "x-api-key: $KEY_B"',
            ),
        ],
        start=1,
    ):
        add_step(doc, n, title_text, body_text, cmd)

    doc.add_heading("8. Database Verification Script", level=1)
    doc.add_paragraph("Use the External DB URL from your laptop. These queries are safe read-only checks.")
    add_code(
        doc,
        'psql "$DATABASE_URL"\n\nselect id, name, plan, status from tenants;\nselect tenant_id, email, role from users;\nselect tenant_id, name, prefix, status from api_keys order by created_at desc;\nselect tenant_id, agent_name, workflow from agent_runs order by created_at desc limit 10;\nselect tenant_id, actor, action, resource, timestamp from audit_logs order by timestamp desc limit 20;\nselect tenant_id, count(*) from agent_runs group by tenant_id;\n\\q',
    )

    doc.add_heading("9. Troubleshooting During Demo", level=1)
    make_table(
        doc,
        ["Symptom", "Likely cause", "Fix"],
        [
            ["localStorage acp_api_key is null", "Expected after login migration", "Use Enterprise Login or generate a tenant API key for Postman"],
            ["User not found", "Placeholder email used", "Query users table and use real owner email"],
            ["Missing audit_logs", "Schema was partially applied", "Run db/schema.sql or migration SQL"],
            ["OpenRouter 429", "Free model rate-limited", "Switch AI_ADVISOR_MODEL to openai/gpt-4o-mini"],
            ["401 from tenant API", "Missing/invalid/revoked acp_ key", "Create a new Tenant API Key"],
            ["Laptop cannot reach DB host", "Internal Render DB URL used", "Use External Database URL locally"],
        ],
        [1.7, 2.2, 2.6],
    )

    doc.add_heading("10. Enterprise Readiness Message", level=1)
    for item in [
        "Human users authenticate with email/password and secure HttpOnly session cookies.",
        "Machine/API traffic uses tenant-scoped acp_ keys that can be created and revoked.",
        "Every core table carries tenant_id and dashboard queries scope through authenticated tenant context.",
        "Audit logs record bootstrap, login, key creation, revocation, connectors, and reset actions.",
        "Next hardening milestone: Postgres Row Level Security and encrypted connector secrets.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    OUT.parent.mkdir(exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
